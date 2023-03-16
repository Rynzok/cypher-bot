from docx import Document


def getting_text_from_a_document(src):
    doc = Document(src)
    message_no_space = ''
    for paragraph in doc.paragraphs:
        message_no_space = "".join(paragraph.text)
    return message_no_space


def writing_text_to_a_document(src, stroka):
    doc = Document(src)
    for paragraph in doc.paragraphs:
        paragraph.text = None
    doc.add_paragraph(stroka)
    doc.save(src)
